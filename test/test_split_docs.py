import unittest
from chain import split_documents, load_documents 
from langchain.text_splitter import CharacterTextSplitter
from docx import Document
import os

# Define the directory path for your test documents
TEST_DIR = "./test/test_documents/"

class TestSplitDocuments(unittest.TestCase):

    def setUp(self):
        # Create a sample document for testing
        document = Document()
        document.add_paragraph("This is a sample document for testing the split_documents function.")

        sample_doc_path = TEST_DIR + 'sample_document.docx'
        document.save(sample_doc_path)
        self.sample_document = load_documents(TEST_DIR)
   
    def tearDown(self):
        # Remove the test files and directories inside TEST_DIR
        for item in os.listdir(TEST_DIR):
            item_path = os.path.join(TEST_DIR, item)
            os.remove(item_path)

    def test_split_documents_with_default_settings(self):
        # Test splitting documents with default settings
        result = split_documents(self.sample_document)
        self.assertTrue(result)
        self.assertEqual(len(result), 1)  # Ensure only one document is returned

    def test_split_documents_with_custom_settings(self):
        # Test splitting documents with custom settings
        result = split_documents(self.sample_document, chunk_size=40, chunk_overlap=0)
        self.assertTrue(result)
        self.assertEqual(len(result), 2)  # Ensure only one document is returned

    def test_split_documents_with_multiple_documents(self):
        # Test splitting multiple documents
        additional_documents = [
            Document(),
            Document(),
        ]
        additional_para = [
            "This is an additional document for testing.",
            "Another additional document with a different structure."
        ]
        for i, doc in enumerate(additional_documents):
            doc.add_paragraph(additional_para[i])
            doc_path = os.path.join(TEST_DIR, f'additional_document_{i}.docx')
            doc.save(doc_path)

        all_documents = load_documents(TEST_DIR)
        result = split_documents(all_documents)
        self.assertTrue(result)
        self.assertEqual(len(result), 3)  # Ensure the number of output documents matches the input